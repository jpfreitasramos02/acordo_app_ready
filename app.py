from flask import Flask, render_template, request, send_file
from io import BytesIO
from docx import Document

app = Flask(__name__)

PERGUNTAS = [
    "1) Como serão votadas as deliberações dos sócios?",
    "2) Como serão designados os administradores da empresa?",
    "3) Qual participação autoriza atos dos administradores previstos no contrato social?",
    "4) Qual participação delibera sobre aumento do capital social?",
    "5) Se houver aumento e o sócio não puder investir, o que acontece?",
    "6) Penalidade ao sócio que contrair dívidas com perdas à Sociedade?",
    "7) Quórum para decidir a remuneração dos administradores?",
    "8) Quórum para decidir remuneração dos sócios? Haverá reinvestimento obrigatório?",
    "9) Haverá capital de giro mínimo? Valor?",
    "10) Venda livre de participações? Oposição à entrada de novos sócios?",
    "11) Quórum para aceitar ingresso de novo sócio (investidor/aceleradora)?",
    "12) Direito de preferência para aquisição de quotas?",
    "13) Tag Along (direito de venda conjunta)?",
    "14) Drag Along (obrigação de venda conjunta)?",
    "15) Direitos de vesting?",
    "16a) Admite sucessão (separação/divórcio/falecimento/falência)?",
    "16b) Como calcular e pagar haveres?",
    "17) Prazo de vigência do Acordo?",
    "18) Não competição / não contratação / não solicitação? Prazo?",
    "19) Informações confidenciais a proteger?",
    "20) Multa por descumprimento da confidencialidade (valor)?",
    "21) Foro comum ou arbitragem? (Câmara, regulamento, composição, sede, língua)"
]

@app.route("/", methods=["GET"])
def index():
    # Renderiza o template corrigido (sem enumerate)
    return render_template("index.html", perguntas=PERGUNTAS)

@app.route("/gerar_docx", methods=["POST"])
def gerar_docx():
    respostas = [request.form.get(f"pergunta_{i}", "") for i in range(len(PERGUNTAS))]

    doc = Document()
    doc.add_heading("Respostas do Questionário — Acordo de Sócios", level=1)

    for pergunta, resposta in zip(PERGUNTAS, respostas):
        doc.add_paragraph(pergunta)
        doc.add_paragraph(f"Resposta: {resposta if resposta else '—'}")

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return send_file(
        buf,
        as_attachment=True,
        download_name="respostas_acordo.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

# Nenhum if __main__ necessário — no Render usaremos gunicorn
